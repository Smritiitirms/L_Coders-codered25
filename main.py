import streamlit as st
import os
import pandas as pd
import plotly.express as px
import sweetviz as sv
from datetime import datetime
import sqlite3
from agents.sanity_check import SanityCheckAgent
from agents.summary import SummaryAgent
from agents.report import ReportAgent

# Set up folders
UPLOAD_FOLDER = "uploads"
REPORT_FOLDER = "reports"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)

sanity_check = SanityCheckAgent()
summary_agent = SummaryAgent()
report_agent = ReportAgent()

# Initialize SQLite database for file history
DB_PATH = "file_history.db"
conn = sqlite3.connect(DB_PATH)
cursor = conn.cursor()

# Create a table for storing file history if not exists
cursor.execute("""
    CREATE TABLE IF NOT EXISTS file_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        file_name TEXT,
        upload_date TEXT
    )
""")
conn.commit()

# Streamlit App Configuration
st.set_page_config(
    page_title="Synapse AI: Medical Data Analysis",
    layout="wide",
    menu_items={"About": "An AI-Powered Medical Data Analysis Tool"},
)

# Sidebar for navigation
st.sidebar.title("🔍 Medical Data Explorer")
page = st.sidebar.radio("Navigation", ["Home", "Feature Analysis", "Diagnostics", "Report", "About"])

# Global variable for loaded data
uploaded_file = None
data = None

# Function to load the dataset
def load_data(file):
    file_type = file.name.split(".")[-1].lower()
    filepath = os.path.join(UPLOAD_FOLDER, file.name)

    with open(filepath, "wb") as f:
        f.write(file.getbuffer())

    if file_type == "csv":
        return pd.read_csv(filepath)
    elif file_type == "xlsx":
        return pd.read_excel(filepath)
    elif file_type == "json":
        return pd.read_json(filepath)
    elif file_type == "parquet":
        return pd.read_parquet(filepath)
    else:
        st.error("Unsupported file format.")
        return None

# Function to log uploaded file to the database
def log_file_history(file_name):
    upload_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cursor.execute("INSERT INTO file_history (file_name, upload_date) VALUES (?, ?)", (file_name, upload_date))
    conn.commit()

# Function to fetch file history from the database
def get_file_history():
    cursor.execute("SELECT file_name, upload_date FROM file_history ORDER BY upload_date DESC")
    return cursor.fetchall()

if page == "Home":
    st.title("📊 Synapse AI")
    st.subheader("Upload and Explore Medical Datasets")

    uploaded_file = st.file_uploader(
        "Upload a file (CSV, Excel, JSON, Parquet)",
        type=["csv", "xlsx", "json", "parquet"],
        accept_multiple_files=False,
    )

    st.subheader("🗂️ File Upload History")
    history = get_file_history()
    if history:
        for file_name, upload_date in history:
            st.write(f"📄 {file_name} (Uploaded on: {upload_date})")
        if st.button("Clear File Upload History"):
            cursor.execute("DELETE FROM file_history")
            conn.commit()
            st.success("File upload history cleared!")
    else:
        st.info("No files uploaded yet.")

    if st.button("Load Data"):
        if uploaded_file:
            data = load_data(uploaded_file)
            if data is not None:
                st.session_state["data"] = data
                st.success("File uploaded successfully!")
                st.dataframe(data.head())
                log_file_history(uploaded_file.name)
        else:
            st.warning("Please upload a file.")

elif page == "Diagnostics":
    st.title("🔬 AI-Driven Diagnostics")
    st.subheader("Predictive Analysis for Medical Data")

    if "data" in st.session_state and st.session_state["data"] is not None:
        data = st.session_state["data"]

        target_column = st.selectbox("Select Target Column (e.g., Diagnosis/Outcome):", data.columns)
        feature_columns = st.multiselect("Select Features for Prediction:", [col for col in data.columns if col != target_column])

        if st.button("Run Predictive Analysis"):
            if len(feature_columns) > 0:
                # Placeholder: Implement predictive model training/inference
                st.info("Training ML Model...")
                
                # Example: Use Logistic Regression (or any other model)
                from sklearn.model_selection import train_test_split
                from sklearn.linear_model import LogisticRegression
                from sklearn.metrics import classification_report
                
                X = data[feature_columns]
                y = data[target_column]
                
                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
                model = LogisticRegression()
                model.fit(X_train, y_train)

                # Predict and Display Results
                predictions = model.predict(X_test)
                report = classification_report(y_test, predictions, output_dict=True)
                st.write("**Classification Report:**")
                st.json(report)
            else:
                st.warning("Please select features for prediction.")
    else:
        st.warning("Upload data first!")

elif page == "Feature Analysis":
    st.title("📈 Feature Analysis")
    if "data" in st.session_state and st.session_state["data"] is not None:
        data = st.session_state["data"]
        
        # Overview
        st.subheader("🔍 Dataset Overview")
        st.write(f"Number of Rows: {data.shape[0]}")
        st.write(f"Number of Columns: {data.shape[1]}")
        st.dataframe(data.head())

        # Visualization Type Selection
        st.subheader("📊 Smart Visualizations")
        vis_type = st.radio("Choose Visualization Type", ["Distribution", "Correlation Matrix", "Scatter Plot", "Box Plot"])

        if vis_type == "Distribution":
            feature_column = st.selectbox("Select a Feature to Analyze", options=data.columns)
            if pd.api.types.is_numeric_dtype(data[feature_column]):
                fig = px.histogram(data, x=feature_column, title=f"Distribution of {feature_column}")
            else:
                fig = px.bar(
                    data[feature_column].value_counts().reset_index(),
                    x="index",
                    y=feature_column,
                    title=f"Value Counts of {feature_column}",
                    labels={"index": feature_column, feature_column: "Count"}
                )
            st.plotly_chart(fig)

        elif vis_type == "Correlation Matrix":
            st.write("**Correlation Matrix for Numeric Features:**")
            corr_matrix = data.corr()
            fig = px.imshow(corr_matrix, text_auto=True, title="Correlation Matrix", color_continuous_scale="Viridis")
            st.plotly_chart(fig)

        elif vis_type == "Scatter Plot":
            x_column = st.selectbox("Select X-axis Feature", options=data.columns)
            y_column = st.selectbox("Select Y-axis Feature", options=data.columns)
            fig = px.scatter(data, x=x_column, y=y_column, title=f"{y_column} vs {x_column}")
            st.plotly_chart(fig)

        elif vis_type == "Box Plot":
            feature_column = st.selectbox("Select a Feature to Analyze", options=data.columns)
            fig = px.box(data, y=feature_column, title=f"Box Plot of {feature_column}")
            st.plotly_chart(fig)

        # Sweetviz Advanced Report Generation
        st.subheader("📄 Generate Advanced Report")
        if st.button("Generate Sweetviz Report"):
            with st.spinner("Generating report..."):
                report = sv.analyze(data)
                report_path = os.path.join(REPORT_FOLDER, "sweetviz_report.html")
                report.show_html(filepath=report_path, open_browser=False)
            
            with open(report_path, "rb") as file:
                st.download_button("Download Sweetviz Report", file, file_name="sweetviz_report.html")
    else:
        st.warning("No data available! Please upload a dataset on the Home page.")
elif page == "Report":
    st.title("📄 Automated Detailed Report")

    if "data" in st.session_state and st.session_state["data"] is not None:
        data = st.session_state["data"]

        st.subheader("📊 Data Overview")
        st.write(f"**Rows:** {data.shape[0]}")
        st.write(f"**Columns:** {data.shape[1]}")
        st.dataframe(data.head())

        # Generate Summary
        st.subheader("🔍 AI-Generated Summary")
        if st.button("Generate Summary"):
            with st.spinner("Generating summary..."):
                # Call local GPT/LMStudio for summary
                from transformers import pipeline

                summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
                summary = summarizer(
                    f"The dataset has {data.shape[0]} rows and {data.shape[1]} columns. "
                    f"Here's a brief description of its contents: {data.describe(include='all').to_string()}",
                    max_length=200,
                    min_length=50,
                    do_sample=False,
                )[0]["summary_text"]

                st.success("Summary Generated:")
                st.write(summary)

        # Generate Detailed Report
        st.subheader("📑 Generate Detailed Report")
        if st.button("Generate Detailed Report"):
            with st.spinner("Generating detailed report..."):
                from docx import Document

                # Create a Word document for the report
                doc = Document()
                doc.add_heading("Automated Data Analysis Report", level=1)

                # Add overview
                doc.add_heading("1. Data Overview", level=2)
                doc.add_paragraph(f"Number of Rows: {data.shape[0]}")
                doc.add_paragraph(f"Number of Columns: {data.shape[1]}")
                doc.add_paragraph("Sample Data:")
                doc.add_paragraph(data.head().to_string())

                # Add statistical analysis
                doc.add_heading("2. Statistical Analysis", level=2)
                doc.add_paragraph(data.describe(include="all").to_string())

                # Add insights (using GPT/local model)
                doc.add_heading("3. AI-Generated Insights", level=2)
                summary = summarizer(
                    f"The dataset has the following key statistical insights: {data.describe(include='all').to_string()}",
                    max_length=200,
                    min_length=50,
                    do_sample=False,
                )[0]["summary_text"]
                doc.add_paragraph(summary)

                # Save the report
                report_path = os.path.join(REPORT_FOLDER, "detailed_report.docx")
                doc.save(report_path)

                with open(report_path, "rb") as file:
                    st.download_button("Download Detailed Report", file, file_name="detailed_report.docx")
    else:
        st.warning("No data available! Please upload a dataset on the Home page.")

